import os
import json
import requests
import time
import PyPDF2
import re
from html import unescape
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
import tempfile
import traceback

try:
    import docx  # python-docx
    DOCX_AVAILABLE = True
except ImportError:
    docx = None
    DOCX_AVAILABLE = False

try:
    from landingai_ade import LandingAIADE
    ADE_SDK_AVAILABLE = True
except ImportError:
    LandingAIADE = None  # type: ignore
    ADE_SDK_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BeautifulSoup = None  # type: ignore
    BS4_AVAILABLE = False

# Import our vector store module and config
from vector_store import create_vector_store
from config import settings

# Determine whether to use mock ADE or real API
USE_MOCK_ADE = not settings.VISION_AGENT_API_KEY

_ade_client: Optional["LandingAIADE"] = None

# The ADE SDK parse endpoint works well, but extract currently expects multipart uploads
# even when passing markdown strings, which conflicts with the documented JSON payload
# (`markdown` + `schema`) described in the official README.
# Until the SDK adds native JSON support for extract, we fall back to the HTTP client.
ENABLE_ADE_SDK_EXTRACT = False

# Regular expressions for cleaning markup-heavy responses
HTML_TAG_PATTERN = re.compile(r"<[^>]+>")
LINE_BREAK_PATTERN = re.compile(r"<\s*(br|/tr|/table|/p)\s*/?>", re.IGNORECASE)
IMPORTANT_METRIC_KEYWORDS = [
    "total",
    "premium",
    "dividend",
    "balance",
    "fee",
    "asset",
    "liabil",
    "equity",
    "cash",
    "revenue",
    "income",
    "sales",
    "payment",
    "adjustment"
]


def get_ade_client() -> Optional["LandingAIADE"]:
    """
    Lazily instantiate the Landing.AI ADE SDK client if available/configured.
    """
    global _ade_client
    if USE_MOCK_ADE or not ADE_SDK_AVAILABLE:
        return None
    if not settings.VISION_AGENT_API_KEY:
        return None
    if _ade_client is None:
        try:
            _ade_client = LandingAIADE(apikey=settings.VISION_AGENT_API_KEY)
        except Exception as exc:
            print(f"Failed to initialize Landing.AI ADE SDK: {exc}")
            _ade_client = None
    return _ade_client


def serialize_ade_object(obj: Any) -> Dict[str, Any]:
    """
    Normalize SDK objects into plain dictionaries so downstream logic
    (debug saves, schema mapping) can expect identical structures whether
    we call via SDK or raw HTTP.
    """
    if isinstance(obj, dict):
        return obj
    
    for method_name in ("to_dict", "dict", "model_dump"):
        method = getattr(obj, method_name, None)
        if callable(method):
            try:
                return method()
            except Exception:
                continue
    
    json_method = getattr(obj, "json", None)
    if callable(json_method):
        try:
            return json.loads(json_method())
        except Exception:
            pass
    
    # Last resort: attempt to serialize via json dumps
    try:
        return json.loads(json.dumps(obj))
    except Exception:
        raise ValueError("Unable to serialize ADE SDK response into dict")

def process_document(file_path: str) -> Dict[str, Any]:
    """
    Process a financial document using Landing.AI ADE.
    Supports: PDF, DOCX/DOC, HTML/HTM, PNG/JPG/JPEG/TIFF (images).
    
    Args:
        file_path: Path to the uploaded file
        
    Returns:
        A dictionary with the processed document data
    """
    print(f"Processing document: {file_path}")
    
    try:
        # Step 1: Extract raw text based on file type (used for fallback / vector store)
        file_ext = Path(file_path).suffix.lower()
        if file_ext in ('.docx', '.doc'):
            pdf_text = extract_text_from_docx(file_path)
        elif file_ext in ('.html', '.htm'):
            pdf_text = extract_text_from_html(file_path)
        elif file_ext in ('.png', '.jpg', '.jpeg', '.tiff', '.tif', '.webp'):
            pdf_text = ""  # ADE handles image files directly — no pre-extraction needed
        else:
            pdf_text = extract_text_from_pdf(file_path)
        
        # Step 2: Call Landing.AI ADE Parse API - this is the main extraction
        # Landing.AI ADE Parse dynamically detects document structure (tables, text, metadata)
        # across ALL pages of the document. We build our data structure from whatever 
        # Landing.AI finds, not a fixed schema. Landing.AI returns chunks from all pages.
        if USE_MOCK_ADE:
            print("Using mock ADE processing")
            ade_response = mock_ade_processing(file_path, pdf_text)
        else:
            print("Calling Landing.AI ADE Parse API (processing all pages)")
            ade_response = call_landing_ai_ade_parse(file_path)
        
        # Step 3: Build financial data structure dynamically from Landing.AI parse response
        # This extracts whatever structure Landing.AI detected (tables, chunks, metadata)
        financial_data = map_to_financial_schema(ade_response, pdf_text)
        
        # ── Build detected_chunks + cell sub-chunks BEFORE vector store ──
        # This must happen before create_vector_store so cell-level vectors
        # are included in the embeddings.
        ade_grounding: dict = {}
        raw_grounding = ade_response.get("grounding")
        if isinstance(raw_grounding, dict):
            ade_grounding = raw_grounding

        _TD_RE = re.compile(r'<td\s+id=["\']([^"\']+)["\']>(.*?)</td>', re.DOTALL | re.IGNORECASE)

        if ade_response.get("chunks"):
            financial_data["detected_chunks"] = []
            for chunk in ade_response["chunks"]:
                chunk_id   = chunk.get("id") or ""
                chunk_type = chunk.get("type") or "text"
                chunk_grounding = chunk.get("grounding", {}) or {}
                chunk_page = chunk_grounding.get("page")
                chunk_box  = chunk_grounding.get("box")

                raw_markdown = chunk.get("markdown") or chunk.get("text") or chunk.get("content") or ""
                cleaned_text = normalize_chunk_text(raw_markdown)

                financial_data["detected_chunks"].append({
                    "id":       chunk_id,
                    "type":     chunk_type,
                    "text":     cleaned_text,
                    "markdown": raw_markdown,
                    "page":     chunk_page,
                    "box":      chunk_box
                })

                if chunk_type == "table" and ade_grounding:
                    for td_id, td_raw_text in _TD_RE.findall(raw_markdown):
                        cell_text = re.sub(r'<[^>]+>', '', td_raw_text).strip()
                        if not cell_text:
                            continue

                        g_entry = ade_grounding.get(td_id)
                        if not g_entry:
                            continue

                        cell_box  = g_entry.get("box")
                        cell_page = g_entry.get("page", chunk_page)
                        pos       = g_entry.get("position", {}) or {}
                        row_idx   = pos.get("row", 0)
                        col_idx   = pos.get("col", 0)

                        financial_data["detected_chunks"].append({
                            "id":              td_id,
                            "type":            "table_cell",
                            "parent_table_id": chunk_id,
                            "text":            cell_text,
                            "markdown":        cell_text,
                            "page":            cell_page,
                            "box":             cell_box,
                            "row":             row_idx,
                            "col":             col_idx,
                        })

                elif chunk_type in ("figure", "chart", "graph"):
                    financial_data["detected_chunks"][-1].update({
                        "subtype":    chunk.get("subtype", chunk_type),
                        "description": chunk.get("description", ""),
                        "chart_data": chunk.get("data", {}),
                    })

        # Generate summary and store supplementary data
        summary = generate_summary(financial_data, pdf_text)
        financial_data["summary"] = summary
        financial_data["original_text"] = pdf_text
        if "markdown" in ade_response:
            financial_data["markdown"] = ade_response["markdown"]
        if ade_grounding:
            financial_data["ade_grounding"] = ade_grounding

        # ── Now create vector store (detected_chunks are available) ──
        document_id = os.path.basename(file_path).split('.')[0]
        vector_store_path = f"./data/vector_stores/{document_id}"
        os.makedirs(vector_store_path, exist_ok=True)

        financial_data_path = os.path.join(vector_store_path, "financial_data.json")
        with open(financial_data_path, "w") as f:
            json.dump(financial_data, f)

        try:
            create_vector_store(financial_data, pdf_text, vector_store_path)
        except Exception as e:
            print(f"Warning: Vector store creation failed (non-critical): {str(e)}")

        # Save debug info
        save_debug_info(document_id, ade_response, {}, financial_data)

        return financial_data
    except Exception as e:
        print(f"Error in process_document: {str(e)}")
        traceback.print_exc()
        
        # Return minimal financial data
        return {
            "metadata": {
                "company_name": "Error Processing Document",
                "document_date": "Unknown",
                "document_type": "Unknown"
            },
            "income_statement": {},
            "balance_sheet": {},
            "cash_flow": {},
            "key_metrics": [],
            "summary": f"Error processing document: {str(e)}"
        }

def extract_text_from_pdf(file_path: str) -> str:
    """Extract plain text from a PDF file using PyPDF2."""
    text = ""
    try:
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        text = ""
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Extract plain text from a DOCX/DOC file using python-docx."""
    if not DOCX_AVAILABLE:
        print("python-docx not installed. Run: pip install python-docx")
        return ""
    try:
        document = docx.Document(file_path)
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        # Also extract table cell text
        for table in document.tables:
            for row in table.rows:
                row_text = "  |  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        return ""


def extract_text_from_html(file_path: str) -> str:
    """Extract plain text from an HTML file using BeautifulSoup."""
    if not BS4_AVAILABLE:
        print("BeautifulSoup not installed. Run: pip install beautifulsoup4")
        return ""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            html_content = f.read()
        soup = BeautifulSoup(html_content, "html.parser")
        # Remove script/style noise
        for tag in soup(["script", "style", "head", "meta"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    except Exception as e:
        print(f"Error extracting text from HTML: {str(e)}")
        return ""

def call_landing_ai_ade_parse(file_path: str) -> Dict[str, Any]:
    """
    Call Landing.AI ADE Parse API to process a document
    """
    client = get_ade_client()
    if client:
        try:
            parse_result = client.parse(
                document=Path(file_path),
                model="dpt-2"
            )
            ade_response = serialize_ade_object(parse_result)
            
            # Save the raw ADE response for debugging
            debug_dir = os.path.join(settings.EXTRACTED_DATA_PATH, os.path.basename(file_path).split('.')[0])
            os.makedirs(debug_dir, exist_ok=True)
            debug_path = os.path.join(debug_dir, "ade_parse_response.json")
            
            with open(debug_path, 'w') as f:
                json.dump(ade_response, f, indent=2)
            
            return ade_response
        except Exception as sdk_error:
            print(f"Landing.AI ADE SDK parse failed, falling back to HTTP: {sdk_error}")
    
    return _call_landing_ai_ade_parse_http(file_path)


def _call_landing_ai_ade_parse_http(file_path: str) -> Dict[str, Any]:
    """
    HTTP fallback for Landing.AI ADE Parse API.
    """
    try:
        # Prepare API key and endpoint
        api_key = settings.VISION_AGENT_API_KEY
        url = f"{settings.ADE_ENDPOINT}/parse"
        
        if not api_key:
            raise Exception("Landing.AI API key not found. Please set VISION_AGENT_API_KEY in your .env file.")
        
        # Prepare headers with authorization
        headers = {
            "Authorization": f"Bearer {api_key}"
        }
        
        # Prepare the file for upload
        with open(file_path, "rb") as file:
            files = {
                "document": (os.path.basename(file_path), file, "application/pdf")
            }
            
            # Optional parameters
            data = {
                "model": "dpt-2-latest"  # Use the latest Landing.AI document model
            }
            
            # Make the API request
            print(f"Sending request to {url}")
            response = requests.post(url, headers=headers, files=files, data=data)
            
            # Check if request was successful
            if response.status_code != 200:
                raise Exception(f"ADE Parse API request failed: {response.text}")
            
            # Parse the response
            ade_response = response.json()
            
            # Save the raw ADE response for debugging
            debug_dir = os.path.join(settings.EXTRACTED_DATA_PATH, os.path.basename(file_path).split('.')[0])
            os.makedirs(debug_dir, exist_ok=True)
            debug_path = os.path.join(debug_dir, "ade_parse_response.json")
            
            with open(debug_path, 'w') as f:
                json.dump(ade_response, f, indent=2)
            
            return ade_response
    
    except Exception as e:
        print(f"Error calling Landing.AI ADE Parse API: {str(e)}")
        # Fall back to mock processing
        return mock_ade_processing(file_path, extract_text_from_pdf(file_path))

def call_landing_ai_ade_extract(markdown_content: str) -> Dict[str, Any]:
    """
    [DEPRECATED] Call Landing.AI ADE Extract API to extract structured data from document
    
    NOTE: This function is no longer used. We now use only Landing.AI ADE Parse API
    which dynamically detects document structure. The structure is built from whatever
    Landing.AI finds in the document, not from a fixed schema.
    
    Uses SDK first, falls back to HTTP if SDK fails.
    """
    # Validate markdown payload early
    if not markdown_content or not markdown_content.strip():
        print("Warning: Empty markdown content, skipping extract")
        return {}
    
    # Try SDK first if enabled
    client = get_ade_client() if ENABLE_ADE_SDK_EXTRACT else None
    if client:
        try:
            financial_schema = _build_financial_schema()
            
            # SDK extract expects markdown (string) and schema (dict)
            extract_result = client.extract(
                markdown=markdown_content,
                schema=financial_schema
            )
            
            # Serialize the response
            if hasattr(extract_result, 'data') and hasattr(extract_result.data, 'extracted_schema'):
                extracted_schema = extract_result.data.extracted_schema
                if isinstance(extracted_schema, dict):
                    return extracted_schema
                # If it's a Pydantic model, convert to dict
                if hasattr(extracted_schema, 'model_dump'):
                    return extracted_schema.model_dump()
                if hasattr(extracted_schema, 'dict'):
                    return extracted_schema.dict()
            
            # Fallback: try to get raw response
            if hasattr(extract_result, 'model_dump'):
                result_dict = extract_result.model_dump()
                if 'data' in result_dict and 'extracted_schema' in result_dict['data']:
                    return result_dict['data']['extracted_schema']
            
            print("Warning: Could not parse SDK extract response, falling back to HTTP")
        except Exception as sdk_error:
            print(f"Landing.AI ADE SDK extract failed, falling back to HTTP: {sdk_error}")
    
    # Fall back to HTTP
    return _call_landing_ai_ade_extract_http(markdown_content)


def _build_financial_schema() -> Dict[str, Any]:
    """
    [DEPRECATED] Define the schema for financial extraction (shared between SDK and HTTP calls).
    
    NOTE: This function is no longer used. We now build the structure dynamically
    from Landing.AI ADE Parse response, adapting to whatever structure Landing.AI detects.
    """
    return {
        "type": "object",
        "properties": {
            "document_type": {
                "type": "string",
                "enum": ["invoice", "financial_report", "bank_statement", "fee_receipt", "other"],
                "description": "Type of financial document"
            },
            "company_name": {
                "type": "string",
                "description": "Name of the company or organization"
            },
            "document_date": {
                "type": "string",
                "description": "Date of the document"
            },
            "total_amount": {
                "type": "number",
                "description": "Total amount or sum on the document"
            },
            "currency": {
                "type": "string",
                "description": "Currency used in the document (USD, EUR, etc.)"
            },
            "line_items": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "description": {
                            "type": "string",
                            "description": "Description of the item"
                        },
                        "amount": {
                            "type": "number",
                            "description": "Amount for this item"
                        },
                        "quantity": {
                            "type": "number",
                            "description": "Quantity of this item"
                        }
                    }
                },
                "description": "Individual line items in the document"
            }
        }
    }


def _call_landing_ai_ade_extract_http(markdown_content: str) -> Dict[str, Any]:
    """
    HTTP fallback for Landing.AI ADE Extract API.
    """
    try:
        # Validate markdown payload early to avoid 400s from the API
        if not markdown_content or not markdown_content.strip():
            raise ValueError("ADE Extract requires non-empty markdown content.")
        
        # Prepare API key and endpoint
        api_key = settings.VISION_AGENT_API_KEY
        url = f"{settings.ADE_ENDPOINT}/extract"
        
        if not api_key:
            raise Exception("Landing.AI API key not found. Please set VISION_AGENT_API_KEY in your .env file.")
        
        # Prepare headers with authorization
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        # Define schema for financial extraction
        financial_schema = _build_financial_schema()
        
        # Prepare the payload
        payload = {
            "schema": financial_schema,
            "markdown": markdown_content
        }
        
        # Make the API request
        response = requests.post(url, headers=headers, json=payload)
        
        # Check if request was successful
        if response.status_code != 200:
            raise Exception(f"ADE Extract API request failed: {response.text}")
        
        # Parse the response
        extract_response = response.json()
        
        # Save the raw extract response for debugging
        debug_dir = os.path.join(settings.EXTRACTED_DATA_PATH, "latest_extraction")
        os.makedirs(debug_dir, exist_ok=True)
        debug_path = os.path.join(debug_dir, "ade_extract_response.json")
        
        with open(debug_path, 'w') as f:
            json.dump(extract_response, f, indent=2)
        
        # Return the extracted schema
        if "data" in extract_response and "extracted_schema" in extract_response["data"]:
            return extract_response["data"]["extracted_schema"]
        return {}
    
    except Exception as e:
        print(f"Error calling Landing.AI ADE Extract API: {str(e)}")
        return {}

def mock_ade_processing(file_path: str, pdf_text: str) -> Dict[str, Any]:
    """
    Mock the Landing.AI ADE API for MVP development
    This creates a simplified version of what ADE would return
    """
    print("Creating mock ADE response")
    
    # Simple text analysis to identify potential tables and sections
    lines = pdf_text.split('\n')
    
    # Find potential financial statement sections
    income_statement_start = -1
    balance_sheet_start = -1
    cash_flow_start = -1
    
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if "income statement" in line_lower or "statement of operations" in line_lower:
            income_statement_start = i
        elif "balance sheet" in line_lower or "statement of financial position" in line_lower:
            balance_sheet_start = i
        elif "cash flow" in line_lower or "statement of cash flows" in line_lower:
            cash_flow_start = i
    
    # Extract company name (simplified)
    company_name = "Example Corporation"
    for i in range(min(10, len(lines))):
        if len(lines[i]) > 5 and "inc" in lines[i].lower() or "corp" in lines[i].lower():
            company_name = lines[i].strip()
            break
    
    # Create mock ADE response
    mock_response = {
        "markdown": pdf_text,
        "chunks": [],
        "metadata": {
            "filename": os.path.basename(file_path),
            "company_name": company_name,
            "document_date": "2023-12-31",  # Placeholder
            "document_type": "Financial Report",
            "page_count": 1,
            "job_id": f"mock-job-{os.path.basename(file_path).split('.')[0]}"
        }
    }
    
    # Add text chunks
    chunk_id = 0
    for i, line in enumerate(lines):
        if len(line.strip()) > 0:
            chunk = {
                "id": f"chunk-{chunk_id}",
                "markdown": line,
                "type": "text",
                "grounding": {
                    "page": 0,
                    "box": {
                        "left": 0,
                        "top": i * 20,
                        "right": 500,
                        "bottom": i * 20 + 20
                    }
                }
            }
            mock_response["chunks"].append(chunk)
            chunk_id += 1
    
    # Add table chunks if we found financial statement sections
    if income_statement_start >= 0:
        table_chunk = {
            "id": f"chunk-table-income",
            "markdown": "\n".join(lines[income_statement_start:income_statement_start + 15]),
            "type": "table",
            "grounding": {
                "page": 0,
                "box": {
                    "left": 0,
                    "top": income_statement_start * 20,
                    "right": 500,
                    "bottom": (income_statement_start + 15) * 20
                }
            }
        }
        mock_response["chunks"].append(table_chunk)
    
    if balance_sheet_start >= 0:
        table_chunk = {
            "id": f"chunk-table-balance",
            "markdown": "\n".join(lines[balance_sheet_start:balance_sheet_start + 15]),
            "type": "table",
            "grounding": {
                "page": 0,
                "box": {
                    "left": 0,
                    "top": balance_sheet_start * 20,
                    "right": 500,
                    "bottom": (balance_sheet_start + 15) * 20
                }
            }
        }
        mock_response["chunks"].append(table_chunk)
    
    if cash_flow_start >= 0:
        table_chunk = {
            "id": f"chunk-table-cashflow",
            "markdown": "\n".join(lines[cash_flow_start:cash_flow_start + 15]),
            "type": "table",
            "grounding": {
                "page": 0,
                "box": {
                    "left": 0,
                    "top": cash_flow_start * 20,
                    "right": 500,
                    "bottom": (cash_flow_start + 15) * 20
                }
            }
        }
        mock_response["chunks"].append(table_chunk)
    
    return mock_response

def normalize_chunk_text(text: str) -> str:
    """
    Convert ADE markdown into plain text lines for downstream parsing.
    """
    if not text:
        return ""
    
    # Standardize line endings
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    
    # Insert line breaks where tables and paragraphs end
    normalized = LINE_BREAK_PATTERN.sub("\n", normalized)
    
    # Remove remaining HTML tags
    normalized = HTML_TAG_PATTERN.sub(" ", normalized)
    
    # Decode HTML entities and normalize whitespace
    normalized = unescape(normalized).replace("\xa0", " ")
    normalized = re.sub(r"\n{2,}", "\n", normalized)
    normalized = re.sub(r"[ \t]{2,}", " ", normalized)
    
    return normalized.strip()


def safely_get_text(item: Any) -> str:
    """
    Safely extract text from an item, handling different formats and types,
    and normalize it for downstream parsing.
    """
    if item is None:
        return ""
    
    if isinstance(item, str):
        return normalize_chunk_text(item)
    
    if isinstance(item, dict):
        # Try common text fields in dictionaries
        for key in ["text", "markdown", "content"]:
            if key in item and item[key]:
                text_content = item[key]
                if isinstance(text_content, str):
                    return normalize_chunk_text(text_content)
    
    # If we can't extract text, return empty string
    return ""


_CURRENCY_PATTERNS = [
    (re.compile(r'\b(?:pak\s*rs\.?|pkr|pakistani\s*rupee)', re.I), "PKR"),
    (re.compile(r'\b(?:inr|indian\s*rupee|₹)', re.I), "INR"),
    (re.compile(r'\b(?:eur|euro|€)', re.I), "EUR"),
    (re.compile(r'\b(?:gbp|pound\s*sterling|£)', re.I), "GBP"),
    (re.compile(r'\b(?:jpy|japanese\s*yen|¥)', re.I), "JPY"),
    (re.compile(r'\b(?:aed|dirham)', re.I), "AED"),
    (re.compile(r'\b(?:sar|saudi\s*riyal)', re.I), "SAR"),
    (re.compile(r'\b(?:rs\.?\s)', re.I), "PKR"),
    (re.compile(r'\$|usd|us\s*dollar', re.I), "USD"),
]


def detect_currency(financial_data: Dict[str, Any]) -> str:
    """Detect the document's currency from table headers, chunk text, and metadata."""
    search_texts: List[str] = []
    for tbl in financial_data.get("tables", []):
        search_texts.extend(tbl.get("header") or [])
        search_texts.append(tbl.get("title") or "")
    for chunk in financial_data.get("detected_chunks", []):
        ctype = chunk.get("type", "")
        if ctype in ("table", "marginalia", "text"):
            search_texts.append((chunk.get("text") or "")[:500])

    combined = " ".join(search_texts)
    for pattern, code in _CURRENCY_PATTERNS:
        if pattern.search(combined):
            return code
    return "USD"


def append_metric(
    financial_data: Dict[str, Any],
    name: str,
    value: Optional[float],
    unit: str = "",
    period: str = "Current Period",
    source: Optional[str] = None,
) -> None:
    """Append a metric if it does not already exist.

    If *unit* is empty or 'USD' (legacy default), we auto-detect from the
    document's table headers / chunk text so the LLM formats the currency
    correctly.
    """
    if value is None:
        return
    if any(metric.get("name", "").lower() == name.lower() for metric in financial_data.get("key_metrics", [])):
        return
    if not unit or unit == "USD":
        unit = detect_currency(financial_data)
    metric_entry = {
        "name": name,
        "value": value,
        "unit": unit,
        "period": period,
    }
    if source:
        metric_entry["source"] = source
    financial_data.setdefault("key_metrics", []).append(metric_entry)


def clean_metric_label(label: str) -> str:
    """Remove numbering/punctuation and trim label text."""
    if not label:
        return ""
    cleaned = re.sub(r"^\d+[\.\)]?\s*", "", label.strip())
    cleaned = cleaned.strip(":")
    return cleaned.strip()


def first_numeric_value(row: Dict[str, str], header: List[str]) -> Optional[float]:
    """Return the first numeric value found in a row, preferring right-most columns."""
    if not row:
        return None
    search_order: List[str] = []
    if header:
        search_order.extend(list(header)[::-1])
    for key in row.keys():
        if key not in search_order:
            search_order.append(key)
    for column in search_order:
        cell_value = row.get(column)
        if not cell_value:
            continue
        numeric_value = extract_number_from_line(cell_value)
        if numeric_value is not None:
            return numeric_value
    return None


def parse_table_rows(raw_text: str, normalized_text: str) -> Tuple[List[str], List[Dict[str, str]]]:
    """
    Parse ADE table markup into header/rows using HTML (preferred) or markdown/whitespace fallbacks.
    """
    header: List[str] = []
    rows: List[Dict[str, str]] = []

    if BS4_AVAILABLE and raw_text and "<table" in raw_text.lower():
        try:
            soup = BeautifulSoup(raw_text, "html.parser")
            table = soup.find("table")
            if table:
                # First, check if there's a thead with th elements
                thead = table.find("thead")
                if thead:
                    header_cells = [cell.get_text(" ", strip=True) for cell in thead.find_all(["th", "td"])]
                    if header_cells:
                        header = header_cells
                
                # Process tbody or all tr elements
                tbody = table.find("tbody")
                tr_elements = tbody.find_all("tr") if tbody else table.find_all("tr")
                
                # Check if first row looks like a header.
                # A row qualifies as a header if every cell is either:
                #   - a bare 4-digit year (e.g. "2018", "2019"), OR
                #   - a short text label with no digit characters
                _YEAR_CELL_RE = re.compile(r'^(19|20)\d{2}$')

                def _any_cell_is_year(cells):
                    """True if at least one cell is a bare 4-digit year."""
                    return any(_YEAR_CELL_RE.match(c.strip()) for c in cells)

                def looks_like_header(cells):
                    if not cells or len(cells) < 2:
                        return False
                    for cell in cells:
                        c = cell.strip()
                        if _YEAR_CELL_RE.match(c):
                            continue  # bare year — valid header cell
                        if len(c) > 60:
                            return False
                        if any(ch.isdigit() for ch in c):
                            return False  # non-year digit content — not a header
                    return True
                
                for tr_idx, tr in enumerate(tr_elements):
                    # Skip if this is a header row (has th elements and we already have header)
                    th_cells = tr.find_all("th")
                    if th_cells and not header:
                        # This row contains header cells
                        header = [cell.get_text(" ", strip=True) for cell in th_cells]
                        continue
                    elif th_cells and header:
                        # Header row but we already have header, skip it
                        continue
                    
                    # Regular data row - handle colspan by duplicating cell content
                    cell_elements = tr.find_all(["td", "th"])
                    cells = []
                    for cell in cell_elements:
                        cell_text = cell.get_text(" ", strip=True)
                        colspan = int(cell.get("colspan", 1))
                        # Add the cell text, and if colspan > 1, add empty strings for spanned columns
                        cells.append(cell_text)
                        for _ in range(colspan - 1):
                            cells.append("")  # Empty cell for colspan
                    
                    if not cells or all(not c.strip() for c in cells):
                        continue
                    
                    # If we don't have a header yet, check if first row looks like a header
                    if not header:
                        if tr_idx == 0 and (looks_like_header(cells) or _any_cell_is_year(cells)) and len(cells) >= 2:
                            # First row is the header (pure labels OR year-based like "Note | 2019 | 2018")
                            header = cells
                            continue
                        else:
                            # First row is data - create generic header
                            if len(cells) == 2:
                                header = ["Field", "Value"]
                            else:
                                header = [f"Column {i+1}" for i in range(len(cells))]
                    
                    # Create row dictionary with header keys
                    row_dict: Dict[str, str] = {}
                    for idx, cell in enumerate(cells):
                        column_name = header[idx] if idx < len(header) else f"col_{idx+1}"
                        row_dict[column_name] = cell
                    rows.append(row_dict)
                
                # If we have rows but no header, create generic header
                if rows and not header:
                    # Use first row keys as header
                    first_row = rows[0] if rows else {}
                    header = list(first_row.keys())
                
                if header or rows:
                    # Debug: print table structure
                    print(f"Parsed table: {len(header)} columns, {len(rows)} rows")
                    print(f"Header: {header}")
                    if rows:
                        print(f"First row keys: {list(rows[0].keys())}")
                    return header, rows
        except Exception as e:
            print(f"Error parsing HTML table: {str(e)}")
            import traceback
            traceback.print_exc()
            pass

    normalized_lines = [line.strip() for line in normalized_text.split('\n') if line.strip()]

    markdown_lines = [line for line in normalized_lines if '|' in line]
    if len(markdown_lines) >= 2:
        parsed_rows: List[List[str]] = []
        for line in markdown_lines:
            stripped = line.strip('|').strip()
            if not stripped:
                continue
            if set(stripped) <= set("-:"):
                continue
            parsed_rows.append([cell.strip() for cell in line.strip('|').split('|')])
        if len(parsed_rows) >= 2:
            header = parsed_rows[0]
            for row_cells in parsed_rows[1:]:
                row_dict = {}
                for idx, cell in enumerate(row_cells):
                    column_name = header[idx] if idx < len(header) else f"col_{idx+1}"
                    row_dict[column_name] = cell
                rows.append(row_dict)
            if rows:
                return header, rows

    fallback_header: List[str] = []
    for line in normalized_lines:
        parts = [segment.strip() for segment in re.split(r"\s{2,}", line) if segment.strip()]
        if len(parts) < 2:
            continue
        if not fallback_header:
            fallback_header = parts
            continue
        row_dict = {}
        for idx, part in enumerate(parts):
            column_name = fallback_header[idx] if idx < len(fallback_header) else f"col_{idx+1}"
            row_dict[column_name] = part
        rows.append(row_dict)
    if rows:
        return fallback_header, rows

    return [], []


def infer_table_title(raw_text: str) -> str:
    """Best-effort extraction of a table title or descriptor from chunk text."""
    normalized = normalize_chunk_text(raw_text)
    lines = normalized.split('\n')
    
    # Look for common table identifiers in the first few lines
    title_keywords = [
        "student copy", "university copy", "bank copy", "fee bill", 
        "invoice", "receipt", "statement", "challan", "form"
    ]
    
    # Check first 10 lines for title-like text
    for line in lines[:10]:
        stripped = line.strip()
        if not stripped:
            continue
        # Skip markdown table separators
        if '|' in stripped and set(stripped.replace('|', '').strip()) <= set('-:'):
            continue
        # Skip numbered lists
        if re.match(r"^\d+[\.\)]", stripped):
            continue
        # Skip if it's just a number or currency
        if re.match(r"^[\d\s,\$€£¥]+$", stripped):
            continue
        # If line contains title keywords or is reasonably long, use it
        if any(keyword in stripped.lower() for keyword in title_keywords) or len(stripped) > 5:
            # Clean up common prefixes
            cleaned = re.sub(r"^(copy|form|bill|statement|invoice|receipt)[\s:]+", "", stripped, flags=re.IGNORECASE)
            return cleaned[:120] if cleaned else stripped[:120]
    
    # Fallback: return first non-empty, non-table line
    for line in lines[:15]:
        stripped = line.strip()
        if stripped and '|' not in stripped and not re.match(r"^[\d\s,\$€£¥\-:]+$", stripped):
            return stripped[:120]
    
    return "Table"


def extract_metrics_from_table_entry(table_entry: Dict[str, Any], financial_data: Dict[str, Any]) -> None:
    """Add key metrics from general financial tables."""
    header = table_entry.get("header") or []
    rows = table_entry.get("rows") or []
    source = table_entry.get("title") or table_entry.get("id")
    
    for row in rows:
        label = ""
        if header:
            label = row.get(header[0], "") or ""
        if not label and row:
            first_key = next(iter(row))
            label = row.get(first_key, "")
        cleaned_label = clean_metric_label(label)
        if not cleaned_label:
            continue
        label_lower = cleaned_label.lower()
        if not any(keyword in label_lower for keyword in IMPORTANT_METRIC_KEYWORDS):
            continue
        value = first_numeric_value(row, header)
        if value is None:
            continue
        
        if "asset" in label_lower and "total" in label_lower:
            financial_data["balance_sheet"]["total_assets"] = value
            append_metric(financial_data, "Total Assets", value, source=source)
        elif "liabil" in label_lower and "total" in label_lower:
            financial_data["balance_sheet"]["total_liabilities"] = value
            append_metric(financial_data, "Total Liabilities", value, source=source)
        elif "equity" in label_lower and "total" in label_lower:
            financial_data["balance_sheet"]["total_equity"] = value
            append_metric(financial_data, "Total Equity", value, source=source)
        elif "cash" in label_lower and ("flow" in label_lower or "operating" in label_lower):
            financial_data["cash_flow"]["operating_cash_flow"] = value
            append_metric(financial_data, "Operating Cash Flow", value, source=source)
        elif "revenue" in label_lower or "sales" in label_lower or "income" in label_lower:
            financial_data["income_statement"].setdefault("line_items", {})[cleaned_label] = value
            append_metric(financial_data, cleaned_label, value, source=source)
        else:
            append_metric(financial_data, cleaned_label, value, source=source)

def map_to_financial_schema(ade_response: Dict[str, Any], pdf_text: str) -> Dict[str, Any]:
    """
    Build financial data structure dynamically from Landing.AI ADE Parse response
    
    This function processes the output from Landing.AI ADE Parse API and extracts
    whatever structure Landing.AI detected (tables, chunks, metadata).
    We don't force a fixed schema - we adapt to what Landing.AI finds in the document.
    """
    try:
        # Initialize flexible financial schema - structure will be built from what Landing.AI finds
        financial_data = {
            "metadata": {
                "company_name": "Unknown Company",
                "document_date": "Unknown Date",
                "document_type": "Document"  # Will be inferred from content
            },
            "income_statement": {},
            "balance_sheet": {},
            "cash_flow": {},
            "key_metrics": [],
            "tables": []
        }
        
        # Extract metadata from ADE response (this is what Landing.AI detected)
        if "metadata" in ade_response:
            metadata = ade_response["metadata"]
            
            # Use metadata if not already set
            if financial_data["metadata"]["company_name"] == "Unknown Company" and "company_name" in metadata:
                financial_data["metadata"]["company_name"] = metadata["company_name"]
            
            if financial_data["metadata"]["document_date"] == "Unknown Date" and "document_date" in metadata:
                financial_data["metadata"]["document_date"] = metadata["document_date"]
            
            # Use document_type from metadata if available
            if "document_type" in metadata:
                financial_data["metadata"]["document_type"] = metadata["document_type"]
        
        # Process chunks to extract financial information
        # Landing.AI Parse returns chunks with types (text, table, etc.) and grounding info
        # We process whatever chunks Landing.AI detected, building our structure dynamically
        if "chunks" in ade_response:
            chunks = ade_response["chunks"]
            
            # ── First pass: extract metadata from chunks (context-aware) ──
            # Helper: reject lines that look like disclaimers / sentences
            def _looks_like_entity(text: str) -> bool:
                t = text.strip()
                if not t or len(t) < 3 or len(t) > 80:
                    return False
                t_lower = t.lower()
                if t_lower.startswith(("the ", "each ", "this ", "these ", "as noted")):
                    return False
                if " include " in t_lower or " intended " in t_lower or " prepared " in t_lower:
                    return False
                if t.count(" ") > 12:
                    return False
                return True

            # Helper: strip IDs / reference numbers from an entity string
            _ID_STRIP_RE = re.compile(r'\s*[\*]+\s*\d+\s*[\*]*')
            def _clean_entity(text: str) -> str:
                cleaned = _ID_STRIP_RE.sub('', text).strip(' *')
                return cleaned if len(cleaned) >= 3 else text.strip()

            # Pass 1a: company from marginalia (title-like short phrases)
            _TITLE_SEP_RE = re.compile(r'\s*[-–—:]\s*')
            _GENERIC_DOC_LABELS = {"challan", "challan form", "invoice", "receipt",
                                   "bill", "fee bill", "statement", "report",
                                   "form", "voucher", "slip", "memo", "notice"}
            for chunk in chunks:
                if financial_data["metadata"]["company_name"] != "Unknown Company":
                    break
                chunk_text = safely_get_text(chunk)
                chunk_type = chunk.get("type", "")
                if not chunk_text:
                    continue
                if chunk_type == "marginalia":
                    for line in chunk_text.split('\n'):
                        line = line.strip()
                        if not _looks_like_entity(line):
                            continue
                        parts = _TITLE_SEP_RE.split(line)
                        non_generic = [p.strip() for p in parts
                                       if p.strip().lower() not in _GENERIC_DOC_LABELS
                                       and len(p.strip()) >= 3
                                       and _looks_like_entity(p.strip())
                                       and not re.match(r'^[\d/\s:,\.aAmMpP]+$', p.strip())]
                        if non_generic:
                            financial_data["metadata"]["company_name"] = _clean_entity(non_generic[0])
                            break
                        for part in parts:
                            part = part.strip()
                            if len(part) >= 3 and _looks_like_entity(part):
                                if not re.match(r'^[\d/\s:,\.aAmMpP]+$', part):
                                    financial_data["metadata"]["company_name"] = _clean_entity(part)
                                    break
                        if financial_data["metadata"]["company_name"] != "Unknown Company":
                            break

            # Pass 1b: company from text/table chunks (US-style suffix, tightened)
            if financial_data["metadata"]["company_name"] == "Unknown Company":
                for chunk in chunks:
                    chunk_text = safely_get_text(chunk)
                    if not chunk_text:
                        continue
                    chunk_text_lower = chunk_text.lower()
                    if any(term in chunk_text_lower for term in ["company", "corporation", "inc", "bank", "limited"]):
                        for line in chunk_text.split('\n'):
                            line = line.strip()
                            if _looks_like_entity(line) and any(t in line.lower() for t in ["inc", "corp", "ltd", "llc", "bank", "limited", "university"]):
                                financial_data["metadata"]["company_name"] = _clean_entity(line)
                                break
                    if financial_data["metadata"]["company_name"] != "Unknown Company":
                        break

            # Pass 1c: date from chunks (quarter/period/ended style)
            for chunk in chunks:
                if financial_data["metadata"]["document_date"] != "Unknown Date":
                    break
                chunk_text = safely_get_text(chunk)
                if not chunk_text:
                    continue
                chunk_text_lower = chunk_text.lower()
                if "date" in chunk_text_lower or "period" in chunk_text_lower or "quarter" in chunk_text_lower:
                    for line in chunk_text.split('\n'):
                        ll = line.lower()
                        if ("quarter" in ll or "period" in ll) and ("ended" in ll or "ending" in ll):
                            financial_data["metadata"]["document_date"] = line.strip()
                            break
            
            # Second pass: Process ALL chunks from Landing.AI (tables, text, charts, etc.)
            # Landing.AI detects structure across the entire document - we capture everything
            # IMPORTANT: Check ALL chunks for tables, regardless of type, because Landing.AI
            # may embed tables in markdown of any chunk type
            processed_table_ids = set()  # Track processed tables to avoid duplicates
            
            for chunk in chunks:
                raw_chunk_text = (
                    chunk.get("markdown")
                    or chunk.get("text")
                    or chunk.get("content")
                    or ""
                )
                chunk_text = safely_get_text(chunk)
                chunk_type = chunk.get("type", "").lower()
                chunk_id = chunk.get("id", "")
                
                # Skip if no valid text
                if not chunk_text and not raw_chunk_text:
                    continue
                    
                chunk_text_lower = chunk_text.lower() if chunk_text else ""
                
                # GENERALIZED TABLE DETECTION: Check for HTML tables in ANY chunk type
                # Landing.AI often embeds tables as HTML in markdown, regardless of chunk type
                has_html_table = "<table" in raw_chunk_text.lower() if raw_chunk_text else False
                has_markdown_table = "|" in raw_chunk_text if raw_chunk_text else False
                looks_like_table = chunk_type == "table" or has_html_table or has_markdown_table
                
                if looks_like_table:
                    # Try to parse as table - this works for HTML tables, markdown tables, or structured text
                    try:
                        header, rows = parse_table_rows(raw_chunk_text, chunk_text)
                    except Exception as e:
                        print(f"Error parsing table from chunk {chunk_id}: {str(e)}")
                        header, rows = [], []
                    
                    # Add table if we have ANY structure (header or rows)
                    # Even partial tables are valuable
                    if header or rows:
                        # Create unique table ID if not provided
                        table_id = chunk_id or f"table-{len(financial_data['tables'])}"
                        
                        # Skip if we've already processed this table
                        if table_id in processed_table_ids:
                            continue
                        processed_table_ids.add(table_id)
                        
                        table_entry = {
                            "id": table_id,
                            "title": infer_table_title(raw_chunk_text or chunk_text),
                            "header": header,
                            "rows": rows,
                            "page": chunk.get("grounding", {}).get("page"),
                            "box": chunk.get("grounding", {}).get("box"),
                            "type": chunk_type or "table"
                        }
                        
                        financial_data["tables"].append(table_entry)
                        print(f"✓ Detected table: {table_entry['title']} with {len(header)} columns, {len(rows)} rows")
                        
                        if rows:  # Only extract metrics if we have rows
                            extract_metrics_from_table_entry(table_entry, financial_data)
                    
                    # Check for financial statement types (for categorization, not filtering)
                    if any(term in chunk_text_lower for term in ["income statement", "statement of operations", "profit and loss"]):
                        extract_income_statement(chunk_text, financial_data)
                    elif any(term in chunk_text_lower for term in ["balance sheet", "financial position", "assets", "liabilities"]):
                        extract_balance_sheet(chunk_text, financial_data)
                    elif any(term in chunk_text_lower for term in ["cash flow", "statement of cash flows", "operating activities"]):
                        extract_cash_flow(chunk_text, financial_data)
            
            # Also extract tables from the full document markdown (if available)
            # Landing.AI may return tables in the main markdown that aren't in individual chunks
            if "markdown" in ade_response and ade_response["markdown"]:
                full_markdown = ade_response["markdown"]
                # Extract all HTML tables from full markdown
                if BS4_AVAILABLE and "<table" in full_markdown.lower():
                    try:
                        soup = BeautifulSoup(full_markdown, "html.parser")
                        all_tables = soup.find_all("table")
                        for table_idx, html_table in enumerate(all_tables):
                            table_html = str(html_table)
                            table_text = html_table.get_text(" ", strip=True)
                            
                            header, rows = parse_table_rows(table_html, table_text)
                            
                            if header or rows:
                                table_id = f"full-markdown-table-{table_idx}"
                                if table_id not in processed_table_ids:
                                    processed_table_ids.add(table_id)
                                    table_entry = {
                                        "id": table_id,
                                        "title": infer_table_title(table_text),
                                        "header": header,
                                        "rows": rows,
                                        "page": None,  # Full markdown doesn't have page info
                                        "box": None,
                                        "type": "markdown_table"
                                    }
                                    financial_data["tables"].append(table_entry)
                                    print(f"✓ Detected table from full markdown: {table_entry['title']} with {len(header)} columns, {len(rows)} rows")
                                    if rows:
                                        extract_metrics_from_table_entry(table_entry, financial_data)
                    except Exception as e:
                        print(f"Error extracting tables from full markdown: {str(e)}")
            
            # Summary: Log how many tables were detected
            total_tables = len(financial_data["tables"])
            if total_tables > 0:
                print(f"✓ Successfully detected {total_tables} table(s) from document")
                for idx, table in enumerate(financial_data["tables"], 1):
                    print(f"  Table {idx}: '{table['title']}' - {len(table.get('header', []))} cols, {len(table.get('rows', []))} rows")
            else:
                print("⚠ Warning: No tables detected in document. This might indicate a parsing issue.")

            # ── Post-table pass: derive company from first table row if still unknown ──
            if financial_data["metadata"]["company_name"] == "Unknown Company" and financial_data["tables"]:
                for tbl in financial_data["tables"]:
                    rows = tbl.get("rows") or []
                    header = tbl.get("header") or []
                    if not rows:
                        continue
                    first_row = rows[0]
                    candidate = ""
                    if header:
                        candidate = str(first_row.get(header[0], "")).strip()
                        if not candidate and len(header) > 1:
                            candidate = str(first_row.get(header[1], "")).strip()
                    else:
                        candidate = str(next(iter(first_row.values()), "")).strip()
                    candidate = _clean_entity(candidate)
                    if candidate and _looks_like_entity(candidate) and not re.match(r'^[\d\s,\.\-]+$', candidate):
                        financial_data["metadata"]["company_name"] = candidate
                        break

            # ── Post-table pass: derive document date from table rows ──
            _DATE_LABELS = {"due date", "date", "document date", "as at", "period ended",
                            "year ended", "report date", "statement date", "effective date"}
            if financial_data["metadata"]["document_date"] == "Unknown Date" and financial_data["tables"]:
                for tbl in financial_data["tables"]:
                    rows = tbl.get("rows") or []
                    header = tbl.get("header") or []
                    for row in rows:
                        label = ""
                        if header:
                            label = str(row.get(header[0], "")).strip()
                        else:
                            label = str(next(iter(row.values()), "")).strip()
                        if label.lower().rstrip(":") in _DATE_LABELS:
                            value = ""
                            if header and len(header) > 1:
                                value = str(row.get(header[1], "")).strip()
                            else:
                                vals = list(row.values())
                                value = str(vals[1]).strip() if len(vals) > 1 else ""
                            if value and len(value) > 4:
                                financial_data["metadata"]["document_date"] = value
                                break
                    if financial_data["metadata"]["document_date"] != "Unknown Date":
                        break
        
        # If no financial statements were found through tables, try regex extraction
        if not financial_data["income_statement"] and not financial_data["balance_sheet"]:
            # Extract revenue
            revenue = extract_metric(pdf_text, ["revenue", "net sales", "total revenue"])
            if revenue:
                append_metric(financial_data, "Revenue", revenue)
                financial_data["income_statement"]["revenue"] = revenue
            
            # Extract net income
            net_income = extract_metric(pdf_text, ["net income", "net earnings", "net profit"])
            if net_income:
                append_metric(financial_data, "Net Income", net_income)
                financial_data["income_statement"]["net_income"] = net_income
            
            # Extract total assets
            total_assets = extract_metric(pdf_text, ["total assets"])
            if total_assets:
                append_metric(financial_data, "Total Assets", total_assets)
                financial_data["balance_sheet"]["total_assets"] = total_assets
            
            # Extract cash flow from operations
            operating_cash = extract_metric(pdf_text, ["cash flow from operations", "operating cash flow", "net cash provided by operating activities"])
            if operating_cash:
                append_metric(financial_data, "Operating Cash Flow", operating_cash)
                financial_data["cash_flow"]["operating_cash_flow"] = operating_cash
        
        # ── Fee amount fallback: only if no table-derived total exists ──
        _has_total_metric = any(
            "total" in m["name"].lower() and ("amount" in m["name"].lower() or "due" in m["name"].lower())
            for m in financial_data["key_metrics"]
        )
        if not _has_total_metric:
            fee_amount = extract_metric(
                pdf_text,
                ["total fee", "total amount", "total due", "amount due", "total amount due"],
                require_total_context=True
            )
            if fee_amount:
                try:
                    numeric_val = float(str(fee_amount).replace(",", "").replace("Rs.", "").replace("Rs", "").strip())
                    if numeric_val < 1e7:
                        if not any(m["name"].lower() == "fee amount" for m in financial_data["key_metrics"]):
                            append_metric(financial_data, "Fee Amount", fee_amount)
                except (ValueError, TypeError):
                    pass
        
        # ── Infer document type: chunk-based first, then pdf_lower fallback ──
        if financial_data["metadata"]["document_type"] == "Document":
            _doc_type_resolved = False

            # Priority 1: check chunk content for high-confidence phrases
            if "chunks" in ade_response:
                for chunk in ade_response["chunks"]:
                    ct = safely_get_text(chunk)
                    if not ct:
                        continue
                    cl = ct.lower()
                    if "financial statement" in cl or "model financial statements" in cl:
                        financial_data["metadata"]["document_type"] = "Financial Statement"
                        _doc_type_resolved = True
                        break
                    if "balance sheet" in cl or "statement of financial position" in cl:
                        financial_data["metadata"]["document_type"] = "Balance Sheet"
                        _doc_type_resolved = True
                        break
                    if "income statement" in cl or "profit and loss" in cl or "statement of operations" in cl:
                        financial_data["metadata"]["document_type"] = "Income Statement"
                        _doc_type_resolved = True
                        break
                    if "cash flow" in cl or "statement of cash flows" in cl:
                        financial_data["metadata"]["document_type"] = "Cash Flow Statement"
                        _doc_type_resolved = True
                        break
                    if "invoice" in cl and chunk.get("type", "") in ("text", "marginalia"):
                        financial_data["metadata"]["document_type"] = "Invoice"
                        _doc_type_resolved = True
                        break

            # Priority 2: pdf_text fallback (financial-statements checked first)
            if not _doc_type_resolved:
                pdf_lower = pdf_text.lower()
                if "financial statement" in pdf_lower or ("statement" in pdf_lower and "financial" in pdf_lower):
                    financial_data["metadata"]["document_type"] = "Financial Statement"
                elif "balance sheet" in pdf_lower:
                    financial_data["metadata"]["document_type"] = "Balance Sheet"
                elif "income statement" in pdf_lower or "profit and loss" in pdf_lower:
                    financial_data["metadata"]["document_type"] = "Income Statement"
                elif "cash flow" in pdf_lower:
                    financial_data["metadata"]["document_type"] = "Cash Flow Statement"
                elif "invoice" in pdf_lower:
                    financial_data["metadata"]["document_type"] = "Invoice"
                elif "receipt" in pdf_lower:
                    financial_data["metadata"]["document_type"] = "Receipt"
                elif "challan" in pdf_lower or ("fee" in pdf_lower and "payment" in pdf_lower):
                    financial_data["metadata"]["document_type"] = "Fee Document"
                elif "fee" in pdf_lower or "payment" in pdf_lower:
                    financial_data["metadata"]["document_type"] = "Fee Document"
                elif "report" in pdf_lower:
                    financial_data["metadata"]["document_type"] = "Financial Report"
        
        return financial_data
    
    except Exception as e:
        print(f"Error in map_to_financial_schema: {str(e)}")
        traceback.print_exc()
        
        # Return minimal financial data
        return {
            "metadata": {
                "company_name": "Error Mapping Schema",
                "document_date": "Unknown",
                "document_type": "Unknown"
            },
            "income_statement": {},
            "balance_sheet": {},
            "cash_flow": {},
            "key_metrics": [],
            "summary": f"Error mapping financial schema: {str(e)}"
        }

def extract_income_statement(table_text: str, financial_data: Dict[str, Any]) -> None:
    """Extract data from income statement table"""
    try:
        lines = table_text.split('\n')
        
        for line in lines:
            line_lower = line.lower()
            
            # Look for revenue
            if "revenue" in line_lower or "sales" in line_lower:
                revenue = extract_number_from_line(line)
                if revenue is not None:
                    financial_data["income_statement"]["revenue"] = revenue
                    append_metric(financial_data, "Revenue", revenue)
            
            # Look for net income
            elif "net income" in line_lower or "net profit" in line_lower or "net earnings" in line_lower:
                net_income = extract_number_from_line(line)
                if net_income is not None:
                    financial_data["income_statement"]["net_income"] = net_income
                    append_metric(financial_data, "Net Income", net_income)
            
            # Look for gross profit
            elif "gross profit" in line_lower or "gross margin" in line_lower:
                gross_profit = extract_number_from_line(line)
                if gross_profit is not None:
                    financial_data["income_statement"]["gross_profit"] = gross_profit
                    append_metric(financial_data, "Gross Profit", gross_profit)
    except Exception as e:
        print(f"Error extracting income statement: {str(e)}")

def extract_balance_sheet(table_text: str, financial_data: Dict[str, Any]) -> None:
    """Extract data from balance sheet table"""
    try:
        lines = table_text.split('\n')
        
        for line in lines:
            line_lower = line.lower()
            
            # Look for total assets
            if "total assets" in line_lower:
                total_assets = extract_number_from_line(line)
                if total_assets is not None:
                    financial_data["balance_sheet"]["total_assets"] = total_assets
                    append_metric(financial_data, "Total Assets", total_assets)
            
            # Look for total liabilities
            elif "total liabilities" in line_lower:
                total_liabilities = extract_number_from_line(line)
                if total_liabilities is not None:
                    financial_data["balance_sheet"]["total_liabilities"] = total_liabilities
                    append_metric(financial_data, "Total Liabilities", total_liabilities)
            
            # Look for equity
            elif "total equity" in line_lower or "shareholders' equity" in line_lower:
                total_equity = extract_number_from_line(line)
                if total_equity is not None:
                    financial_data["balance_sheet"]["total_equity"] = total_equity
                    append_metric(financial_data, "Total Equity", total_equity)
    except Exception as e:
        print(f"Error extracting balance sheet: {str(e)}")

def extract_cash_flow(table_text: str, financial_data: Dict[str, Any]) -> None:
    """Extract data from cash flow statement table"""
    try:
        lines = table_text.split('\n')
        
        for line in lines:
            line_lower = line.lower()
            
            # Look for operating cash flow
            if "operating activities" in line_lower or "cash flow from operations" in line_lower:
                operating_cash = extract_number_from_line(line)
                if operating_cash is not None:
                    financial_data["cash_flow"]["operating_cash_flow"] = operating_cash
                    append_metric(financial_data, "Operating Cash Flow", operating_cash)
            
            # Look for investing cash flow
            elif "investing activities" in line_lower or "cash flow from investing" in line_lower:
                investing_cash = extract_number_from_line(line)
                if investing_cash is not None:
                    financial_data["cash_flow"]["investing_cash_flow"] = investing_cash
                    append_metric(financial_data, "Investing Cash Flow", investing_cash)
            
            # Look for financing cash flow
            elif "financing activities" in line_lower or "cash flow from financing" in line_lower:
                financing_cash = extract_number_from_line(line)
                if financing_cash is not None:
                    financial_data["cash_flow"]["financing_cash_flow"] = financing_cash
                    append_metric(financial_data, "Financing Cash Flow", financing_cash)
    except Exception as e:
        print(f"Error extracting cash flow: {str(e)}")

def extract_number_from_line(line: str) -> Optional[float]:
    """Extract numeric value from a line of text"""
    try:
        if not line:
            return None
        multiplier = 1
        lower_line = line.lower()
        if re.search(r"\bcr\b", lower_line):
            multiplier = -1
        elif re.search(r"\bdr\b", lower_line):
            multiplier = 1
        
        # Remove common characters that might interfere with number parsing
        cleaned = (
            line.replace(',', '')
            .replace('$', '')
            .replace('(', '-')
            .replace(')', '')
        )
        cleaned = re.sub(r"\b(dr|cr)\b", "", cleaned, flags=re.IGNORECASE)
        
        # Find all potential numbers in the line
        words = cleaned.split()
        
        for word in words:
            try:
                # Try to convert to float
                value = float(word)
                # Skip small values like years
                if abs(value) >= 100:  # Assume values less than 100 are not financial figures
                    return multiplier * value
            except ValueError:
                continue
        
        # If no suitable number found, return None
        return None
    except Exception:
        return None

def extract_metric(text: str, keywords: List[str], *, require_total_context: bool = False) -> Optional[float]:
    """
    Extract a financial metric from text based on keywords.

    When *require_total_context* is True the matching line must also contain
    'total' or 'due' (or equivalent context words).  This prevents picking up
    stray numbers (like late-fee rates or student IDs) when scanning pdf_text.
    """
    _CONTEXT_WORDS = {"total", "due", "net", "payable", "balance"}
    try:
        lines = text.lower().split('\n')

        for line in lines:
            for keyword in keywords:
                if keyword in line:
                    if require_total_context and not any(cw in line for cw in _CONTEXT_WORDS):
                        continue
                    number = extract_number_from_line(line)
                    if number is not None:
                        return number

        return None
    except Exception:
        return None

def generate_summary(financial_data: Dict[str, Any], pdf_text: str) -> str:
    """
    Generate a simple summary of the financial document
    """
    try:
        company = financial_data.get("metadata", {}).get("company_name", "The company")
        document_type = financial_data.get("metadata", {}).get("document_type", "Financial Report")
        document_date = financial_data.get("metadata", {}).get("document_date", "")
        
        summary = f"Financial report for {company}.\n\n"
        
        if document_date:
            summary = f"{document_type} for {company} dated {document_date}.\n\n"
        
        # Add key metrics to summary
        metrics = financial_data.get("key_metrics", [])
        if metrics:
            summary += "Key financial metrics:\n"
            for metric in metrics:
                name = metric.get("name", "")
                value = metric.get("value", "")
                unit = metric.get("unit", "")
                
                if name and value:
                    # Format the value for display
                    if isinstance(value, (int, float)):
                        # Format large numbers with commas
                        formatted_value = f"${value:,.2f}" if unit == "USD" else f"{value:,}"
                    else:
                        formatted_value = value
                    
                    summary += f"- {name}: {formatted_value}\n"
        
        return summary
    except Exception as e:
        print(f"Error generating summary: {str(e)}")
        return "Error generating summary."

def save_debug_info(document_id: str, ade_response: Dict[str, Any], extracted_data: Dict[str, Any], financial_data: Dict[str, Any]) -> None:
    """
    Save debug information for troubleshooting
    """
    try:
        debug_dir = os.path.join(settings.EXTRACTED_DATA_PATH, document_id)
        os.makedirs(debug_dir, exist_ok=True)
        
        # Save financial data
        financial_data_path = os.path.join(debug_dir, "financial_data.json")
        with open(financial_data_path, 'w') as f:
            json.dump(financial_data, f, indent=2)
        
        # Save ADE response if not already saved
        ade_response_path = os.path.join(debug_dir, "ade_response.json")
        if not os.path.exists(ade_response_path):
            with open(ade_response_path, 'w') as f:
                json.dump(ade_response, f, indent=2)
        
        # Save extracted data if available
        if extracted_data:
            extracted_data_path = os.path.join(debug_dir, "extracted_data.json")
            with open(extracted_data_path, 'w') as f:
                json.dump(extracted_data, f, indent=2)
    except Exception as e:
        print(f"Error saving debug info: {str(e)}")